import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_to_doc(doc, table_name, table_description, columns_df):
    doc.add_heading(table_name, level=1)
    doc.add_paragraph(table_description)

    # Add a table for the columns
    columns_table = doc.add_table(rows=1, cols=6)

    # Define the header row
    hdr_cells = columns_table.rows[0].cells
    hdr_cells[0].text = "Column Name"
    hdr_cells[1].text = "Data Type"
    hdr_cells[2].text = "Constraint"
    hdr_cells[3].text = "Description"
    hdr_cells[4].text = "Example 1"
    hdr_cells[5].text = "Example 2"

    # Add the column details to the table
    table_columns = columns_df[columns_df['Table'] == table_name]
    for _, row in table_columns.iterrows():
        row_cells = columns_table.add_row().cells
        row_cells[0].text = str(row['Column Name'])
        row_cells[1].text = str(row['Data Type'])
        row_cells[2].text = str(row['Constraint'])
        row_cells[3].text = str(row['Description'])
        row_cells[4].text = str(row['Example 1'])
        row_cells[5].text = str(row['Example 2'])

    # Apply border style to the table
    for row in columns_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def generate_documentation(excel_path, output_path):
    # Load the Excel file
    excel_data = pd.ExcelFile(excel_path)

    # Load the data from each sheet
    tables_df = pd.read_excel(excel_path, sheet_name='Tables')
    columns_df = pd.read_excel(excel_path, sheet_name='Columns')

    # Create a new Document
    doc = Document()

    # Iterate through each table in the Tables sheet
    for _, row in tables_df.iterrows():
        add_table_to_doc(doc, row['Table'], row['Description'], columns_df)

    # Save the document
    doc.save(output_path)

if __name__ == "__main__":
    excel_path = "../docs/Original Tables Documentation.xlsx"
    output_path = "../docs/all_tables_documentation_with_borders.docx"
    generate_documentation(excel_path, output_path)
